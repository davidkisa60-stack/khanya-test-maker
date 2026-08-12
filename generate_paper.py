#!/usr/bin/env python3
"""
Enhanced generate_paper.py for Khanya Test Maker
- Clean titles and no branding
- Sequential numbering
- No dotted dashes (replaced by real working space)
- Improved math display: unicode superscripts for indices, bolder fraction bars
- Extra spacing between sub-parts (a), (b), (i), etc.
- Supports PDF (reportlab) and Word .docx (python-docx)
- Custom paper title supported
"""

import argparse
import json
import random
import re
from pathlib import Path
from io import BytesIO

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm, mm
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, HRFlowable
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

try:
    pdfmetrics.registerFont(TTFont('DejaVu', '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf'))
    pdfmetrics.registerFont(TTFont('DejaVuBold', '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'))
    MAIN_FONT = 'DejaVu'
    BOLD_FONT = 'DejaVuBold'
except:
    MAIN_FONT = 'Helvetica'
    BOLD_FONT = 'Helvetica-Bold'

# Optional Word support
try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

BASE_DIR = Path(__file__).parent
IMAGES_DIR = BASE_DIR / "assets" / "images"

# Per-subject data paths (matches JS getSubjectDataPath)
SUBJECT_DATA_MAP = {
    "Mathematics": BASE_DIR / "subjects" / "mathematics" / "data" / "questions.json",
    "Biology": BASE_DIR / "subjects" / "biology" / "data" / "questions.json",
    "Physical Science": BASE_DIR / "subjects" / "physical_science" / "data" / "questions.json",
    "Physics": BASE_DIR / "subjects" / "physics" / "data" / "questions.json",
    "Economics": BASE_DIR / "subjects" / "economics" / "data" / "questions.json",
    "Chemistry": BASE_DIR / "subjects" / "chemistry" / "data" / "questions.json",
    "Development Studies": BASE_DIR / "subjects" / "development_studies" / "data" / "questions.json",
    "Accounting": BASE_DIR / "subjects" / "accounting" / "data" / "questions.json",
}

def get_subject_data_path(subject="Mathematics"):
    """Return the questions.json path for the given subject (case sensitive as in UI)."""
    key = subject or "Mathematics"
    return SUBJECT_DATA_MAP.get(key, SUBJECT_DATA_MAP["Mathematics"])


def load_questions(subject="Mathematics"):
    """Load questions for a specific subject (supports per-subject data)."""
    path = get_subject_data_path(subject)
    if not path.exists():
        print(f"Warning: No questions file for {subject} at {path}. Using empty list.")
        return []
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)




def get_question_by_id(qid, questions):
    for q in questions:
        if q["id"] == qid:
            return q
    return None


def pretty_math(text):
    """Improve display of indices, matrices, and fraction bars."""
    if not text:
        return text

    # Unicode superscripts for powers/indices (0-9)
    sup_map = str.maketrans("0123456789", "⁰¹²³⁴⁵⁶⁷⁸⁹")
    text = re.sub(r'\^(\d+)', lambda m: m.group(1).translate(sup_map), text)

    # Make fraction division lines look bolder (text representation)
    text = text.replace('────────────', '════════════════════')
    text = text.replace('──────', '══════════')
    text = text.replace('────────', '════════════')

    # Minor matrix improvement for readability (keeps it simple)
    text = re.sub(r'\((\d+)\s+(\d+)\s*;\s*(\d+)\s+(\d+)\)', r'⎛\1  \2⎞\n⎝\3  \4⎠', text)

    return text


def clean_body_for_pdf(body):
    """Remove dotted answer lines and prepare clean text for display."""
    if not body:
        return ""
    cleaned = body
    cleaned = re.sub(r'\.{3,}', '', cleaned)          # remove all long dots/dashes
    cleaned = re.sub(r'\s*\[\d+\]', '', cleaned)      # remove [n] marks
    cleaned = cleaned.replace('$', '').replace('\\times', '×').replace('\\pi', 'π')
    cleaned = cleaned.replace('\\', '')
    cleaned = re.sub(r' +', ' ', cleaned)
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned).strip()
    cleaned = pretty_math(cleaned)
    return cleaned


def build_latex_document(selected_questions, title="Test Paper"):
    """Optional LaTeX output (unchanged structure)."""
    latex = r"""\documentclass[11pt,a4paper]{article}
\usepackage[margin=1.5cm]{geometry}
\usepackage{amsmath,amssymb}
\usepackage{enumitem}
\usepackage{graphicx}
\usepackage{fancyhdr}
\pagestyle{fancy}
\fancyhf{}
\rhead{\small }
\lhead{\small }
\title{\textbf{""" + title + r"""}}
\date{}
\begin{document}
\maketitle
\noindent\textbf{Instructions:} Answer all questions. Show your working where required.

\vspace{0.5cm}
"""
    for idx, q in enumerate(selected_questions, 1):
        latex += f"\n\\noindent\\textbf{{{idx}. {q['title']}}} \\hfill [{q['total_marks']} marks]\n\n"
        latex += q['latex'] + "\n\n"
        for img in q.get("images", []):
            latex += f"\\begin{{center}}\\includegraphics[width=0.65\\textwidth]{{assets/images/{img}}}\\end{{center}}\n\\vspace{{0.2cm}}\n"
        latex += "\\vspace{0.4cm}\n"
    latex += r"\end{document}"
    return latex



def add_page_number(canvas, doc):
    """Draw page number at the top center of every page."""
    canvas.saveState()
    canvas.setFont('Helvetica', 9)
    page_num = canvas.getPageNumber()
    text = f"Page {page_num}"
    canvas.drawCentredString(A4[0] / 2, A4[1] - 1.0*cm, text)
    canvas.restoreState()

def build_pdf(selected_questions, title="Test Paper"):
    """PDF with improved math, spacing between parts, working space, and images.
       No title, no instructions. Page numbers at top.
    """
    buf = BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        rightMargin=1.5*cm, leftMargin=1.5*cm,
        topMargin=1.8*cm, bottomMargin=1.4*cm,
        onFirstPage=add_page_number,
        onLaterPages=add_page_number
    )
    styles = getSampleStyleSheet()
    q_style = ParagraphStyle('Q', fontName=BOLD_FONT, fontSize=11, spaceBefore=10, spaceAfter=4)
    body_style = ParagraphStyle('Body', fontName=MAIN_FONT, fontSize=10, leading=14, spaceAfter=2, alignment=TA_LEFT)

    story = []
    # No title, no instructions (user request). Page numbers added via onFirstPage/onLaterPages
    story.append(Spacer(1, 2*mm))

    for idx, q in enumerate(selected_questions, 1):
        story.append(Paragraph(f"<b>{idx}.</b>  <font color='#15803d'>[{q['total_marks']} marks]</font>", q_style))

        body = clean_body_for_pdf(q.get('body_markdown', q.get('latex', '')))

        # Split body and add extra space between sub-parts (a), (b), (i), etc.
        parts = re.split(r'(?=\([a-z]\)|[ \t]*\([ivx]+\))', body)
        for part in parts:
            part = part.strip()
            if not part:
                continue
            para_text = part.replace('\n', '<br/>')
            story.append(Paragraph(para_text, body_style))

            # Extra breathing room between sub-questions
            if re.search(r'\([a-z]\)|[ivx]+\)', part):
                story.append(Spacer(1, 5*mm))

        # Working space (generous, not dashes)
        story.append(Spacer(1, 12*mm))

        # Images (only attached ones, properly scaled)
        for img_name in q.get("images", []):
            p = IMAGES_DIR / img_name
            if p.exists():
                try:
                    story.append(Spacer(1, 2*mm))
                    img = Image(str(p))
                    max_w = 11 * cm
                    max_h = 8 * cm
                    if img.imageWidth > 0 and img.imageHeight > 0:
                        scale = min(max_w / img.imageWidth, max_h / img.imageHeight)
                        img.drawWidth = img.imageWidth * scale
                        img.drawHeight = img.imageHeight * scale
                    else:
                        img.drawWidth = max_w
                        img.drawHeight = max_h
                    story.append(img)
                    story.append(Spacer(1, 2*mm))
                except Exception as e:
                    print(f"Warning: could not add image {img_name}: {e}")
                    pass

        story.append(Spacer(1, 6*mm))

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


def build_docx(selected_questions, title="Test Paper"):
    """Word document export with improved formatting."""
    if not HAS_DOCX:
        raise ImportError("python-docx is not installed. Please run: pip install python-docx")

    doc = Document()

    # Set nice margins
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.4)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

    # Title
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # No instructions (per request)
    doc.add_paragraph()  # breathing space

    for idx, q in enumerate(selected_questions, 1):
        # Question header
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}.")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(f"   [{q['total_marks']} marks]")
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(21, 128, 61)   # emerald green

        # Body text with improvements
        body = clean_body_for_pdf(q.get('body_markdown', q.get('latex', '')))

        # Add paragraphs with extra space after sub-parts
        for line in body.split('\n'):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(3)

            if re.search(r'\([a-z]\)|[ivx]+\)', line):
                p.paragraph_format.space_after = Pt(8)   # extra space between (a), (b), etc.

        # Working space (blank paragraphs)
        for _ in range(3):
            doc.add_paragraph()

        # Images
        for img_name in q.get("images", []):
            pth = IMAGES_DIR / img_name
            if pth.exists():
                try:
                    doc.add_picture(str(pth), width=Inches(4.5))
                    doc.add_paragraph()
                except Exception as e:
                    print(f"Warning: could not add image to docx {img_name}: {e}")

        doc.add_paragraph()  # extra gap before next question

    # Return as bytes
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--ids")
    parser.add_argument("--random", action="store_true")
    parser.add_argument("--num", type=int, default=6)
    parser.add_argument("--output", default="paper")
    parser.add_argument("--latex-only", action="store_true")
    parser.add_argument("--title", default="Test Paper")
    parser.add_argument("--subject", default="Mathematics",
                        help="Subject name e.g. Mathematics, Biology, Physics (matches UI)")
    parser.add_argument("--format", choices=["pdf", "docx"], default="pdf",
                        help="Output format: pdf or docx (default: pdf)")
    args = parser.parse_args()

    qs = load_questions(args.subject)
    sel = []
    if args.ids:
        for i in args.ids.split(","):
            q = get_question_by_id(i.strip(), qs)
            if q:
                sel.append(q)
    elif args.random:
        sel = random.sample(qs, min(args.num, len(qs)))

    if not sel:
        print("No questions selected")
        return

    # Always write LaTeX (for reference)
    tex = build_latex_document(sel, title=args.title)
    Path(args.output).with_suffix(".tex").write_text(tex, encoding="utf-8")
    print("LaTeX written")

    if args.format == "docx":
        if not HAS_DOCX:
            print("Error: python-docx not installed. Run: pip install python-docx")
            return
        docx_bytes = build_docx(sel, title=args.title)
        out_path = Path(args.output).with_suffix(".docx")
        out_path.write_bytes(docx_bytes)
        print("Word document written")
    else:
        # PDF (default)
        pdfb = build_pdf(sel, title=args.title)
        Path(args.output).with_suffix(".pdf").write_bytes(pdfb)
        print("PDF written")

    print("Done with", len(sel), "questions")


if __name__ == "__main__":
    main()
